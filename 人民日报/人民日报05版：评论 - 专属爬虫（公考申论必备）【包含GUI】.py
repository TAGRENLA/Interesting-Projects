'''
作者：github（tagrenla） WeiXin（S64959）
特别鸣谢：github（caspiankexin）为本项目提供思路来源！！！
'''

import requests
import bs4
import os
import datetime
import time
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox


def fetchUrl(url):
    '''
    功能：访问 url 的网页，获取网页内容并返回
    参数：目标网页的 url
    返回：目标网页的 html 内容
    '''
    headers = {
        'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36',
    }
    r = requests.get(url, headers=headers)
    r.raise_for_status()
    r.encoding = r.apparent_encoding
    return r.text


def getPageList(year, month, day):
    '''
    功能：获取当天报纸的各版面的链接列表
    参数：年，月，日
    '''
    url = f'http://paper.people.com.cn/rmrb/pc/layout/{year}{month}/{day}/node_01.html'
    html = fetchUrl(url)
    bsobj = bs4.BeautifulSoup(html, 'html.parser')
    temp = bsobj.find('div', attrs={'id': 'pageList'})
    if temp:
        pageList = temp.ul.find_all('div', attrs={'class': 'right_title-name'})
    else:
        pageList = bsobj.find('div', attrs={'class': 'swiper-container'}).find_all('div',
                                                                                   attrs={'class': 'swiper-slide'})

    linkList = []
    for page in pageList:
        link = page.a["href"]
        url = f'http://paper.people.com.cn/rmrb/pc/layout/{year}{month}/{day}/{link}'
        if 'node_05' in str(url):
            linkList.append(url)
    return linkList


def getTitleList(year, month, day, pageUrl):
    '''
    功能：获取报纸某一版面的文章链接列表
    参数：年，月，日，该版面的链接
    '''
    html = fetchUrl(pageUrl)
    bsobj = bs4.BeautifulSoup(html, 'html.parser')
    temp = bsobj.find('div', attrs={'id': 'titleList'})
    if temp:
        titleList = temp.ul.find_all('li')
    else:
        titleList = bsobj.find('ul', attrs={'class': 'news-list'}).find_all('li')

    linkList = []
    for title in titleList:
        tempList = title.find_all('a')
        for temp in tempList:
            link = temp["href"]
            if 'content' in link:
                url = f'http://paper.people.com.cn/rmrb/pc/content/{year}{month}/{day}/{link}'
                linkList.append(url)
    return linkList


def getContent(html):
    '''
    功能：解析 HTML 网页，获取新闻的文章内容
    参数：html 网页内容
    '''
    bsobj = bs4.BeautifulSoup(html, 'html.parser')

    # 获取文章 标题
    title = bsobj.h3.text + '\n' + bsobj.h1.text + '\n' + bsobj.h2.text + '\n'

    # 获取文章 内容
    pList = bsobj.find('div', attrs={'id': 'ozoom'}).find_all('p')
    content = ''
    for p in pList:
        content += p.text + '\n'

    # 返回结果 标题+内容
    resp = title + content
    return resp


def save_to_word(content, path, filename):
    '''
    功能：将所有文章内容保存到一个Word文档中，文章间用分割线隔开
    参数：要保存的内容，路径，文件名
    '''
    if not os.path.exists(path):
        os.makedirs(path)

    file_path = os.path.join(path, filename)

    # 如果Word文件已存在，则打开，否则创建新的
    if os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document()

    doc.add_paragraph(content)
    doc.add_paragraph("\n" + "-" * 50 + "\n")  # 分割线

    doc.save(file_path)


def download_rmrb_to_word(year, month, day, destdir):
    '''
    功能：爬取《人民日报》网站某年某月某日的新闻内容，并保存在一个Word文件中
    '''
    pageList = getPageList(year, month, day)

    word_filename = f"{year}{month}{day}.docx"
    save_path = os.path.join(destdir, f"{year}{month}{day}")

    for pageNo, page in enumerate(pageList, start=1):
        try:
            titleList = getTitleList(year, month, day, page)
            for titleNo, url in enumerate(titleList, start=1):
                html = fetchUrl(url)
                content = getContent(html)

                # 保存到Word文件
                save_to_word(content, save_path, word_filename)
        except Exception as e:
            print(f"日期 {year}-{month}-{day} 下的版面 {page} 出现错误：{e}")
            continue


def gen_dates(b_date, days):
    day = datetime.timedelta(days=1)
    for i in range(days):
        yield b_date + day * i


def get_date_list(beginDate, endDate):
    """
    获取日期列表
    :param start: 开始日期
    :param end: 结束日期
    :return: 开始日期和结束日期之间的日期列表
    """

    start = datetime.datetime.strptime(beginDate, "%Y%m%d")
    end = datetime.datetime.strptime(endDate, "%Y%m%d")

    data = []
    for d in gen_dates(start, (end - start).days):
        data.append(d)
    return data


def on_start():
    '''
    当点击"开始爬取"按钮时执行的函数
    '''
    begin_date = begin_date_entry.get()
    end_date = end_date_entry.get()
    dest_dir = dest_dir_entry.get()

    if not begin_date or not end_date or not dest_dir:
        messagebox.showerror("错误", "请确保所有字段都已填写")
        return

    try:
        data = get_date_list(begin_date, end_date)
        for d in data:
            year = str(d.year)
            month = str(d.month) if d.month >= 10 else '0' + str(d.month)
            day = str(d.day) if d.day >= 10 else '0' + str(d.day)

            download_rmrb_to_word(year, month, day, dest_dir)
            progress_label.config(text=f"爬取完成：{year}{month}{day}")
        messagebox.showinfo("完成", "数据爬取完成！")
    except Exception as e:
        messagebox.showerror("错误", f"发生错误：{e}")


def on_browse():
    '''
    打开文件夹选择对话框，选择保存路径
    '''
    folder_selected = filedialog.askdirectory()
    if folder_selected:
        dest_dir_entry.delete(0, tk.END)
        dest_dir_entry.insert(0, folder_selected)


def on_today():
    '''
    填充今天的日期
    '''
    today = datetime.datetime.today()
    begin_date = today.strftime("%Y%m%d")
    end_date = (today + datetime.timedelta(days=1)).strftime("%Y%m%d")
    begin_date_entry.delete(0, tk.END)
    end_date_entry.delete(0, tk.END)
    begin_date_entry.insert(0, begin_date)
    end_date_entry.insert(0, end_date)


def on_yesterday():
    '''
    填充昨天的日期
    '''
    yesterday = datetime.datetime.today() - datetime.timedelta(days=1)
    begin_date = yesterday.strftime("%Y%m%d")
    end_date = (yesterday + datetime.timedelta(days=1)).strftime("%Y%m%d")
    begin_date_entry.delete(0, tk.END)
    end_date_entry.delete(0, tk.END)
    begin_date_entry.insert(0, begin_date)
    end_date_entry.insert(0, end_date)


# 创建主窗口
root = tk.Tk()
root.title("人民日报05版：评论 - 专属爬虫（公考申论必备）")

# 创建输入字段和标签
tk.Label(root, text="开始日期 (格式：YYYYMMDD):").grid(row=0, column=0, padx=10, pady=10)
begin_date_entry = tk.Entry(root, width=15)
begin_date_entry.grid(row=0, column=1)

tk.Label(root, text="结束日期 (格式：YYYYMMDD):").grid(row=1, column=0, padx=10, pady=10)
end_date_entry = tk.Entry(root, width=15)
end_date_entry.grid(row=1, column=1)


tk.Label(root, text="保存路径:").grid(row=2, column=0, padx=10, pady=10)
# 设置默认路径
default_path = r"E:\Users\TAGRENLA\Desktop\rmrb"
dest_dir_entry = tk.Entry(root, width=30)
dest_dir_entry.grid(row=2, column=1)
# 设置输入框的默认值
dest_dir_entry.insert(0, default_path)


browse_button = tk.Button(root, text="浏览...", command=on_browse)
browse_button.grid(row=2, column=2)

# 创建"开始爬取"按钮
start_button = tk.Button(root, text="开始爬取", command=on_start)
start_button.grid(row=3, column=0, columnspan=3, pady=20)

# 创建"今天的报纸"按钮
today_button = tk.Button(root, text="今天的报纸", command=on_today)
today_button.grid(row=4, column=0)

# 创建"昨天的报纸"按钮
yesterday_button = tk.Button(root, text="昨天的报纸", command=on_yesterday)
yesterday_button.grid(row=4, column=1)

# 创建进度标签
progress_label = tk.Label(root, text="")
progress_label.grid(row=5, column=0, columnspan=3)

# 运行主循环
root.mainloop()
