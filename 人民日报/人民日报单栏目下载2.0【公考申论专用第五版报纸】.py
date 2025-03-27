# ecoding:utf-8 
'''
代码名称：爬取人民日报数据为txt文件
编写日期：2025年1月1日
原作者：github（caspiankexin）
更改编辑者：github（TAGRENLA）
版本：第3版
可爬取的时间范围：2024年12月起
注意：此代码仅供交流学习，不得作为其他用途。
'''


import requests
import bs4
import os
import datetime
import time
from docx import Document
import os
def fetchUrl(url):
    '''
    功能：访问 url 的网页，获取网页内容并返回
    参数：目标网页的 url
    返回：目标网页的 html 内容
    '''

    headers = {
        'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36',
    }

    r = requests.get(url,headers=headers)
    r.raise_for_status()
    r.encoding = r.apparent_encoding
    return r.text

def getPageList(year, month, day):
    '''
    功能：获取当天报纸的各版面的链接列表
    参数：年，月，日
    '''
    # if day == version:
    url = 'http://paper.people.com.cn/rmrb/pc/layout/' + year + month + '/' + day + '/node_01.html'
    html = fetchUrl(url)
    bsobj = bs4.BeautifulSoup(html,'html.parser')
    temp = bsobj.find('div', attrs = {'id': 'pageList'})
    if temp:
        pageList = temp.ul.find_all('div', attrs = {'class': 'right_title-name'})
    else:
        pageList = bsobj.find('div', attrs = {'class': 'swiper-container'}).find_all('div', attrs = {'class': 'swiper-slide'})
    linkList = []

    for page in pageList:
        link = page.a["href"]
        url = 'http://paper.people.com.cn/rmrb/pc/layout/'  + year + month + '/' + day + '/' + link
        """
        ⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀
        ⠀⠀⠀⠀⠀⠀⠀⢀⣀⣀⡀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀
        ⠀⠀⠀⠀⢀⣤⡾⠿⠿⠿⠿⢷⡀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀
        ⠀⠀⠀⣰⡿⠁⠀⠀⠀⠀⠀⠈⢷⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀
        ⠀⠀⠀⣿⠁⠀⢀⡀⠀⢀⡀⠀⢸⡇⠀⠀⠀⠀⠀⠀⠀⠀⠀
        ⠀⠀⠀⢿⡄⠀⠈⠃⠀⠈⠃⠀⣸⠇⠀⠀⠀⠀⠀⠀⠀⠀⠀
        ⠀⠀⠀⠈⠻⣦⣀⠀⠀⠀⢀⡼⠁⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀
        ⠀⠀⠀⠀⠀⠈⠙⠛⠛⠛⠉⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀

        ███████╗ █████╗ ██╗     ███████╗
        ██╔════╝██╔══██╗██║     ██╔════╝
        █████╗  ███████║██║     █████╗  
        ██╔══╝  ██╔══██║██║     ██╔══╝  
        ██║     ██║  ██║███████╗███████╗
        ╚═╝     ╚═╝  ╚═╝╚══════╝╚══════╝

        ⚠️  这里是需要修改的地方！请务必替换代码！ ⚠️
        ------------------------------------------------
        同时函数 save_to_word  download_rmrb_to_word 也进行了修改，导出word格式，方便打印
        """
        # ⬇️ 这里是代码需要修改的地方 ⬇️


        if 'node_05' in str(url):   # 这行是代码的精髓
            linkList.append(url)
    return linkList

def getTitleList(year, month, day, pageUrl):
    '''
    功能：获取报纸某一版面的文章链接列表
    参数：年，月，日，该版面的链接
    '''
    html = fetchUrl(pageUrl)
    bsobj = bs4.BeautifulSoup(html,'html.parser')
    temp = bsobj.find('div', attrs = {'id': 'titleList'})
    if temp:
        titleList = temp.ul.find_all('li')
    else:
        titleList = bsobj.find('ul', attrs = {'class': 'news-list'}).find_all('li')
    linkList = []

    for title in titleList:
        tempList = title.find_all('a')
        for temp in tempList:
            link = temp["href"]
            if 'content' in link:
                url = 'http://paper.people.com.cn/rmrb/pc/content/' + year + month + '/' + day + '/' + link
                linkList.append(url)
    return linkList

def getContent(html):
    '''
    功能：解析 HTML 网页，获取新闻的文章内容
    参数：html 网页内容
    '''
    bsobj = bs4.BeautifulSoup(html,'html.parser')

    # 获取文章 标题
    title = bsobj.h3.text + '\n' + bsobj.h1.text + '\n' + bsobj.h2.text + '\n'
    #print(title)

    # 获取文章 内容
    pList = bsobj.find('div', attrs = {'id': 'ozoom'}).find_all('p')
    content = ''
    for p in pList:
        content += p.text + '\n'
    #print(content)

    # 返回结果 标题+内容
    resp = title + content
    return resp




def save_to_word(content, path, filename):
    '''
    功能：将所有文章内容保存到一个Word文档中，文章间用分割线隔开
    参数：要保存的内容，路径，文件名
    '''
    if not os.path.exists(path):
        os.makedirs(path)

    file_path = os.path.join(path, filename)

    # 如果Word文件已存在，则打开，否则创建新的
    if os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document()

    doc.add_paragraph(content)
    doc.add_paragraph("\n" + "-" * 50 + "\n")  # 分割线

    doc.save(file_path)


def download_rmrb_to_word(year, month, day, destdir):
    '''
    功能：爬取《人民日报》网站某年某月某日的新闻内容，并保存在一个Word文件中
    '''
    pageList = getPageList(year, month, day)

    word_filename = f"{year}{month}{day}.docx"
    save_path = os.path.join(destdir, f"{year}{month}{day}")

    for pageNo, page in enumerate(pageList, start=1):
        try:
            titleList = getTitleList(year, month, day, page)
            for titleNo, url in enumerate(titleList, start=1):
                html = fetchUrl(url)
                content = getContent(html)

                # 保存到Word文件
                save_to_word(content, save_path, word_filename)
        except Exception as e:
            print(f"日期 {year}-{month}-{day} 下的版面 {page} 出现错误：{e}")
            continue
def gen_dates(b_date, days):
    day = datetime.timedelta(days = 1)
    for i in range(days):
        yield b_date + day * i


def get_date_list(beginDate, endDate):
    """
    获取日期列表
    :param start: 开始日期
    :param end: 结束日期
    :return: 开始日期和结束日期之间的日期列表
    """

    start = datetime.datetime.strptime(beginDate, "%Y%m%d")
    end = datetime.datetime.strptime(endDate, "%Y%m%d")

    data = []
    for d in gen_dates(start, (end-start).days):
        data.append(d)
    return data


if __name__ == '__main__':
    '''
    主函数：程序入口
    '''
    # 输入起止日期，爬取之间的新闻
    print("欢迎使用人民日报爬虫，请输入以下信息：")
    beginDate = input('请输入开始日期:')
    # beginDate = '20250327'
    # endDate = '20250328'
    endDate = input('请输入结束日期:')
    destdir = input("请输入数据保存的地址：") #  r"E:\Users\TAGRENLA\Desktop\rmrb"
    data = get_date_list(beginDate, endDate)

    for d in data:
        year = str(d.year)
        month = str(d.month) if d.month >=10 else '0' + str(d.month)
        day = str(d.day) if d.day >=10 else '0' + str(d.day)
        destdir = destdir  # 爬下来的文件的存储地方

        download_rmrb_to_word(year, month, day, destdir)
        print("爬取完成：" + year + month + day)
    lastend = input("本月数据爬取完成!可以关闭软件了")

