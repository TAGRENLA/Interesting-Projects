import requests
import bs4
import os
import datetime
import time
from docx import Document

def fetchUrl(url):
    headers = {
        'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36',
    }
    r = requests.get(url, headers=headers)
    r.raise_for_status()
    r.encoding = r.apparent_encoding
    return r.text

def getPageList(year, month, day):
    url = f'http://paper.people.com.cn/rmrb/pc/layout/{year}{month}/{day}/node_01.html'
    html = fetchUrl(url)
    bsobj = bs4.BeautifulSoup(html, 'html.parser')
    temp = bsobj.find('div', attrs={'id': 'pageList'})
    pageList = temp.ul.find_all('div', attrs={'class': 'right_title-name'}) if temp else \
               bsobj.find('div', attrs={'class': 'swiper-container'}).find_all('div', attrs={'class': 'swiper-slide'})
    
    return [f'http://paper.people.com.cn/rmrb/pc/layout/{year}{month}/{day}/{page.a["href"]}' for page in pageList]

def getTitleList(year, month, day, pageUrl):
    html = fetchUrl(pageUrl)
    bsobj = bs4.BeautifulSoup(html, 'html.parser')
    temp = bsobj.find('div', attrs={'id': 'titleList'})
    titleList = temp.ul.find_all('li') if temp else \
                bsobj.find('ul', attrs={'class': 'news-list'}).find_all('li')
    
    return [f'http://paper.people.com.cn/rmrb/pc/content/{year}{month}/{day}/{temp["href"]}' \
            for title in titleList for temp in title.find_all('a') if 'content' in temp["href"]]

def getContent(html):
    bsobj = bs4.BeautifulSoup(html, 'html.parser')
    title = '\n'.join([bsobj.h3.text, bsobj.h1.text, bsobj.h2.text])
    content = '\n'.join(p.text for p in bsobj.find('div', attrs={'id': 'ozoom'}).find_all('p'))
    return title + '\n' + content

def download_rmrb(year, month, day, destdir):
    document = Document()
    pageList = getPageList(year, month, day)
    for pageNo, page in enumerate(pageList, start=1):
        try:
            titleList = getTitleList(year, month, day, page)
            for titleNo, url in enumerate(titleList, start=1):
                html = fetchUrl(url)
                content = getContent(html)
                filename = f'{year}{month}{day}-{str(pageNo).zfill(2)}-{str(titleNo).zfill(2)}'
                document.add_paragraph(f'======={filename}======')
                document.add_paragraph(content)
        except Exception as e:
            print(f"日期 {year}-{month}-{day} 下的版面 {page} 出现错误：{e}")
            continue
    word_filename = f'{year}{month}{day}_rmrb.docx'
    document.save(os.path.join(destdir, word_filename))

def gen_dates(b_date, days):
    return (b_date + datetime.timedelta(days=i) for i in range(days))

def get_date_list(beginDate, endDate):
    start = datetime.datetime.strptime(beginDate, "%Y%m%d")
    end = datetime.datetime.strptime(endDate, "%Y%m%d")
    return list(gen_dates(start, (end - start).days))

if __name__ == '__main__':
    print("欢迎使用人民日报爬虫，请输入以下信息：")
    beginDate = input('请输入开始日期:')
    endDate = input('请输入结束日期:')
    destdir = r"E:\\Users\\TAGRENLA\\Desktop\\rmrb"
    for d in get_date_list(beginDate, endDate):
        year, month, day = str(d.year), f'{d.month:02d}', f'{d.day:02d}'
        download_rmrb(year, month, day, destdir)
        print(f"爬取完成：{year}{month}{day}")
        time.sleep(5)
    input("本月数据爬取完成!可以关闭软件了")
