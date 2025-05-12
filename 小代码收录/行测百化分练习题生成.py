# ecoding:utf-8
# 当前的测试文件名称: b.py
# 作者：TAGRENLA
# 日期：2025-05-12
# 功能：闲来无事，学习行测，编写个代码，联系联系百化分！！！
import random
from fractions import Fraction
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# 数据准备
fractions = [Fraction(1, d) for d in range(2, 20)]
fraction_to_percent = {f: round(float(f) * 100, 2) for f in fractions}

def format_percent(value):
    formatted_number = f"{value:.2f}".rstrip('0').rstrip('.')
    return f"{formatted_number}%"

# 生成题目与答案
def generate_problems(num_problems=80):
    problems, answers = [], []
    for _ in range(num_problems):
        frac = random.choice(fractions)
        percent = format_percent(fraction_to_percent[frac])
        if random.choice([True, False]):
            problems.append(f"{frac} = ________%")
            answers.append(f"{percent}")
        else:
            problems.append(f"{percent} = ________")
            answers.append(f"{frac}")
    return problems, answers

def add_vertical_table(doc, items, with_question=True):
    rows_per_col = len(items) // 4 + (1 if len(items) % 4 else 0)
    table = doc.add_table(rows=rows_per_col, cols=4)
    table.autofit = True

    for col in range(4):
        for row in range(rows_per_col):
            idx = col * rows_per_col + row
            if idx < len(items):
                cell = table.cell(row, col)
                # 统一设置段落格式
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(3)  # 行距
                    paragraph.paragraph_format.line_spacing = 1.5   # 行高
                # 仅保留第一个段落的文本内容
                para = cell.paragraphs[0]
                content = f"({idx+1:>1}) {items[idx]}"  # 带括号的题号
                run = para.add_run(content)
                run.font.size = Pt(11)


# Word 文档生成
def create_word_doc(problems, answers, code, filename="分数百分数互换练习.docx"):
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)

    # 日期与标题
    date_str = datetime.now().strftime("%Y-%m-%d")
    code_str = f"练习编码：{code}"

    # 标题
    title_para = doc.add_paragraph()
    run = title_para.add_run(f"分数与百分数互换练习题（{date_str}）")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 编码
    code_para = doc.add_paragraph(code_str)
    code_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph()
    add_vertical_table(doc, problems, with_question=True)

    # 答案页
    doc.add_page_break()

    answer_title = doc.add_paragraph()
    run = answer_title.add_run("参考答案")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    answer_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    code_para2 = doc.add_paragraph(code_str)
    code_para2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph()
    add_vertical_table(doc, answers, with_question=False)

    doc.save(filename)
    print(f"✅ 练习文档已生成：{filename}（编码：{code}）")

# 主程序
if __name__ == "__main__":
    try:
        count = int(input("请输入题目数量：").strip())
        problems, answers = generate_problems(count)
        today = datetime.now().strftime("%Y%m%d")
        rand_suffix = str(random.randint(1000, 9999))
        code = f"{today}-{rand_suffix}"
        filename = f"练习_{code}.docx"
        create_word_doc(problems, answers, code, filename)
    except ValueError:
        print("❌ 输入无效，请输入整数。")
