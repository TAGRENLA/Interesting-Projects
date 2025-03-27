# encoding:utf-8

'''
禁止外传
更改内容
帖子首页网址
是否只看楼主，输入楼主名称即可
最后页码
登录账号机器密码
文件名

'''


import time
from selenium import webdriver
from selenium.common.exceptions import WebDriverException
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def click_look_at_louzhu(driver):
    print('=============【看楼主】点击开始！=================')
    # if input('是否进行点击按钮【看楼主】？') == 'y':  # 测试用
    if 1:
        print("当前网址:", driver.current_url)
        element = driver.find_element(By.CLASS_NAME, "see-host")
        element.click()
        print(f"成功点击按钮 【看楼主】")
    # a = input('=============【看楼主】点击结束！=================\n')


def close_the_browser(driver):
    """
    提示用户是否关闭浏览器
    :param driver:初始化 Chrome 浏览器
    :return: None
    """
    close_browser = input("是否关闭浏览器？（是/否）: ").strip().lower()
    if close_browser in ['是', 'y', 'yes']:
        driver.quit()
        print("浏览器已关闭。")
    else:
        print("浏览器保持打开状态。")


def open_browser():
    """
    初始化 Chrome 浏览器
    :return: driver
    """
    # 设置 Chrome 浏览器选项
    options = webdriver.ChromeOptions()
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument(
        "user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.5735.199 Safari/537.36")
    # 初始化 Chrome 浏览器
    driver = webdriver.Chrome(options=options)
    return driver



def write_to_word(filename, actor, time, zhengwen):
    """
    将爬取的内容写入 Word 文档（追加模式）。
    :param filename: Word 文件名
    :param actor: 作者（爬取内容）
    :param time: 时间（爬取内容）
    :param zhengwen: 正文（爬取内容）
    """
    # 如果文件不存在，会自动创建
    doc = Document(filename) if os.path.exists(filename) else Document()

    # 设置页边距
    sections = doc.sections
    for section in sections:
        # 页边距（top, bottom, left, right）
        section.top_margin = Pt(20)  # 1英寸 = 72磅
        section.bottom_margin = Pt(20)
        section.left_margin = Pt(20)
        section.right_margin = Pt(20)

    # 写入内容到文档
    doc.add_paragraph(f"楼主：{actor} 时间: {time}")
    doc.add_paragraph(zhengwen)

    # 保存文档
    doc.save(filename)


def pa_chong(driver):
    print('=====================开始进行爬虫操作======================')
    # 爬取帖子时间、作者、正文
    # 先定位到特定父级范围的元素
    parent_div = driver.find_element(By.CLASS_NAME, "atl-main")
    # 在该父级范围内进一步查找子元素
    content_elements_time = driver.find_elements(By.CLASS_NAME, "time")
    content_elements_actor = driver.find_elements(By.CLASS_NAME, "js-vip-check")
    content_elements_zhengwen = driver.find_elements(By.CLASS_NAME, "bbs-content")
    if i > 2:
        content_elements_time.pop(0)
        content_elements_actor.pop(0)
    # print(f'测试当前未第{i-1}页')
    # print(len(content_elements_time),len(content_elements_actor),len(content_elements_zhengwen))
    # 检查是否成功获取到内容
    if len(content_elements_time) == 0 or len(content_elements_actor) == 0 or len(content_elements_zhengwen) == 0:
        print("没有找到符合条件的元素，请检查网页结构或类名是否正确。")
    else:
        # 打印每个帖子的时间、作者、正文内容
        for time, actor, zhengwen in zip(content_elements_time, content_elements_actor, content_elements_zhengwen):
            if actor.text =='目光呆滞的润土':
                write_to_word('风云(中国第一本企业家自传).docx', actor.text, time.text,zhengwen.text)
                print(f'第{i - 1}页  作者：{actor.text}  时间:{time.text} 写入word完毕')
                # print(zhengwen.text)

def next_page(driver):
    print('=============【下一页】点击开始！=================')
    # c = input('是否进行点击按钮【下一页】？')  # 测试的时候用
    c = 'y'
    if c == 'y':
        element = driver.find_element(By.XPATH, f'//a[@onclick="return go({i});"]')
        element.click()
        print(f"成功点击按钮 【下一页】")
    # a = input('==============【下一页】点击结束！==================\n')


def login(driver):
    # 登录功能
    print("正在尝试登录...")
    username = '规则是分层的工具'
    password = '031202'
    # 等待登录表单加载完成
    wait = WebDriverWait(driver, 10)
    username_input = wait.until(EC.presence_of_element_located((By.ID, "vwriter2")))
    password_input = driver.find_element(By.ID, "vpassword")
    # 输入用户名和密码
    username_input.send_keys(username)
    time.sleep(2)
    password_input.send_keys(password)
    time.sleep(2)
    # 点击登录按钮
    login_button = driver.find_element(By.CLASS_NAME, "submit-btn")
    login_button.click()
    print("登录操作已执行，请查看页面状态！\n")

try:
    # 初始化 Chrome 浏览器
    driver = open_browser()
    url = "https://tianya.my/thread/0/11864046/1.html"
    driver.get(url)
    print("当前登录网址:", driver.current_url)
    login(driver)

    for i in range(2,392 + 1):
        print("当前网址:", driver.current_url)

        driver.get(driver.current_url)
        time.sleep(1)
        pa_chong(driver)
        time.sleep(1)
        next_page(driver)
        time.sleep(1)






except WebDriverException as e:
    print(f"WebDriver 错误: {e}")
except Exception as ex:
    print(f"其他错误: {ex}")
finally:
    # 确保资源释放
    if 'driver' in locals() and driver.service.process:
        close_the_browser(driver)
        print("程序已结束，但浏览器可能仍在运行。")
